"""D2R Traderie Tracker 사용자 매뉴얼 Word 파일 생성"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent.parent.parent / "docs" / "tracker" / "user-manual.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# ── 페이지 여백 설정 ────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── 스타일 헬퍼 ────────────────────────────────────────────────────
def set_font(run, size, bold=False, color=None):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "맑은 고딕"
    return p

def body(doc, text, bold_prefix=None):
    """본문. bold_prefix가 있으면 그 부분만 굵게"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, 10, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, 10)
    else:
        r = p.add_run(text)
        set_font(r, 10)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    set_font(r, 10)
    return p

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"※ {text}")
    set_font(r, 9, color=(120, 120, 120))
    return p

def warn(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"⚠ {text}")
    set_font(r, 10, bold=True, color=(200, 60, 60))
    return p

def kv(doc, key, val):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{key}  ")
    set_font(r1, 10, bold=True, color=(60, 140, 60))
    r2 = p.add_run(val)
    set_font(r2, 10)
    return p

def numbered_step(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(f"{num}.  ")
    set_font(r1, 10, bold=True, color=(160, 120, 40))
    r2 = p.add_run(text)
    set_font(r2, 10)
    return p

def add_table(doc, headers, rows, col_widths):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # 헤더
    hdr = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr.cells[i]
        cell.width = Cm(w)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, 10, bold=True, color=(140, 100, 20))
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    # 데이터
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, (val, w) in enumerate(zip(row, col_widths)):
            cell = tr.cells[ci]
            cell.width = Cm(w)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(val)
            set_font(run, 9)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    doc.add_paragraph()

def divider(doc):
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════
# 제목
# ═══════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("D2R Traderie Tracker")
set_font(r, 24, bold=True, color=(140, 100, 20))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
r2 = p2.add_run("사용자 매뉴얼")
set_font(r2, 16, color=(80, 80, 80))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("디아블로2 레저렉션  아이템 시세 자동 조회 프로그램")
set_font(r3, 11, color=(120, 120, 120))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 1. 기본 사용법
# ═══════════════════════════════════════════════════════════════════
heading(doc, "기본 사용법", 1)

heading(doc, "시작 순서", 2)
numbered_step(doc, 1, "최초 실행 시 PaddleOCR 모델 다운로드 화면이 나타납니다.\n'다운로드' 버튼을 눌러 설치하세요. (약 300MB, 최초 1회)")
numbered_step(doc, 2, "'설정' 탭에서 캡처키, 글자색 등 본인 환경에 맞게 설정하세요.")
numbered_step(doc, 3, "디아블로2 레저렉션 실행 후, 화면 상단 프로세스 목록에서 게임을 선택하세요.")
numbered_step(doc, 4, "'실행' 버튼을 누르면 프로그램이 트레이로 최소화됩니다.")
numbered_step(doc, 5, "게임 중 아이템을 획득하면 설정한 캡처키를 눌러주세요.\n시세 정보가 오버레이로 표시됩니다.")

divider(doc)
heading(doc, "무료 / 유료 비교", 2)
add_table(doc,
    headers=["구분",        "무료버전",          "유료버전"          ],
    rows=[
        ["OCR 방식",    "PaddleOCR (로컬)", "AI Vision (클라우드)"],
        ["사용 횟수",   "실행당 20회",      "무제한"              ],
        ["즐겨찾기 갱신","수동만 가능",     "자동 갱신 지원"      ],
        ["오버레이",    "공통 지원",        "공통 지원"           ],
    ],
    col_widths=[4, 6, 6]
)
note(doc, "무료버전도 충분히 사용 가능합니다. OCR 인식률 차이는 크지 않습니다.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 2. 탭별 사용법
# ═══════════════════════════════════════════════════════════════════
heading(doc, "탭별 상세 사용법", 1)

heading(doc, "시세파악 탭", 2)
kv(doc, "시즌",       "래더 / 스탠다드 선택 (복수선택 불가)")
kv(doc, "모드",       "소프트코어 / 하드코어 선택 (복수선택 불가)")
kv(doc, "버전",       "게임 버전 선택 (복수선택 가능)")
kv(doc, "가격시즌",   "가격 기준 시즌 선택 (복수선택 불가)")
kv(doc, "지우기",     "현재 스캔 기록 전체 삭제")
kv(doc, "즐겨찾기",   "즐겨찾기 탭에 추가 (지우기로 삭제해도 즐겨찾기는 유지됨)")
kv(doc, "재조회",     "현재 스캔 목록 시세 재조회")
note(doc, "항목을 더블클릭하면 Traderie 링크를 브라우저로 열어줍니다.")

heading(doc, "즐겨찾기 탭", 2)
bullet(doc, "시세파악 탭에서 '즐겨찾기' 버튼으로 추가한 아이템 목록입니다.")
bullet(doc, "목록 선택 시 오른쪽에서 옵션 수치를 변경할 수 있습니다.")
bullet(doc, "'조회' 버튼으로 시세를 재조회합니다.")
bullet(doc, "항목 더블클릭 시 Traderie 링크를 열어줍니다.")

heading(doc, "설정 탭", 2)
kv(doc, "추출키",       "캡처를 실행할 키 (스킬키와 겹치지 않도록 설정)")
kv(doc, "OCR 설정",     "무료: PaddleOCR  /  유료: AI 방식 선택 가능")
kv(doc, "최대값 여유",  "옵션 검색 범위 설정")
note(doc, "예) 175 방어 + 최대값여유 5  →  175~180 범위로 Traderie 검색")
kv(doc, "로그",         "Traderie 링크 로그 파일 저장 여부 설정")
kv(doc, "데이터",       "GitHub 아이템 목록 갱신 (앱 시작 시 자동 갱신)")
kv(doc, "유료인증",     "유료 키 입력란 (별도 구매)")
kv(doc, "즐겨찾기갱신", "즐겨찾기 가격 자동 갱신 주기 설정 (유료 전용)")
kv(doc, "오버레이",     "가격 표시 위치 / 배경색 / 글자색 / 미리보기 설정")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 3. AI OCR 설정
# ═══════════════════════════════════════════════════════════════════
heading(doc, "AI OCR 설정 (유료버전)", 1)

body(doc, "유료버전에서는 클라우드 AI를 활용한 OCR을 사용할 수 있습니다.")
body(doc, "Groq API는 무료 할당량이 넉넉해 실질적으로 무료로 사용 가능합니다.")
divider(doc)

heading(doc, "Groq API 키 발급 방법", 2)
numbered_step(doc, 1, "구글에서 'groq api 키 발급' 검색 후 console.groq.com 접속")
numbered_step(doc, 2, "Google 계정으로 가입 (새 계정 사용 권장)")
numbered_step(doc, 3, "상단 메뉴에서 'API Keys' → '+Create API Key' 버튼 클릭")
numbered_step(doc, 4, "키 이름 입력  →  Expiration : 'No expiration'  →  Submit")
numbered_step(doc, 5, "발급된 키를 'Copy' 버튼으로 복사 (재조회 불가, 반드시 저장)")
numbered_step(doc, 6, "프로그램 실행 → 유료인증 → OCR 설정을 'AI' 로 변경\n'Groq API' 입력란에 복사한 키 붙여넣기")
note(doc, "키 형태 예시 :  gsk_xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx")
note(doc, "무료 할당량 소진 시 Gemini, GPT 등 다른 AI로 자동 전환됩니다.")

divider(doc)
heading(doc, "AI 제공사 우선순위 변경", 2)
body(doc, "설정 탭 → AI 설정에서 위아래 버튼으로 호출 순서를 변경할 수 있습니다.")
kv(doc, "Groq (Llama 4)", "기본 추천 — 빠르고 무료 할당량 넉넉")
kv(doc, "Gemini",         "Google AI — 안정적")
kv(doc, "GPT-4o",         "OpenAI — 가장 정확하나 유료")
kv(doc, "HuggingFace",    "오픈소스 모델")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 4. 완전 삭제 방법
# ═══════════════════════════════════════════════════════════════════
heading(doc, "완전 삭제 방법", 1)

body(doc, "프로그램을 완전히 제거하려면 아래 항목을 모두 삭제해주세요.")
divider(doc)

add_table(doc,
    headers=["항목",              "경로",                                        "해당 조건"          ],
    rows=[
        ["프로그램 폴더",     "압축 해제한 D2R_Tracker 폴더 전체",          "항상"               ],
        ["설정 파일",         "tracker_config.json  (exe 옆)",               "항상"               ],
        ["API 키 파일",       "tracker_config_keys.json  (exe 옆)",          "AI 키 저장 시"      ],
        ["PaddleOCR 모델",   r"%USERPROFILE%\.paddleocr\  (~300MB)",         "PaddleOCR 사용 시"  ],
        ["로그 파일",         r"내 문서\  (*.txt)",                           "로그 저장 사용 시"  ],
        ["임시 캡처 이미지", r"%TEMP%\d2r_*.png",                            "비정상 종료 시 잔류"],
    ],
    col_widths=[4, 8, 4]
)

heading(doc, "PaddleOCR 모델 폴더 삭제", 2)
body(doc, "Windows 탐색기 주소창에 아래 경로를 입력 후 폴더 삭제:")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
r = p.add_run(r"%USERPROFILE%\.paddleocr")
set_font(r, 10, bold=True, color=(100, 160, 100))

heading(doc, "임시 파일 정리", 2)
body(doc, "Windows 탐색기 주소창에 아래 경로를 입력:")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
r = p.add_run(r"%TEMP%")
set_font(r, 10, bold=True, color=(100, 160, 100))
note(doc, "'d2r_' 로 시작하는 .png 파일을 검색 후 삭제하면 됩니다.")

divider(doc)
warn(doc, "설정 파일(tracker_config.json)에는 AI API 키가 저장되어 있습니다.")
warn(doc, "PC를 다른 사람과 공유한다면 반드시 삭제해주세요.")


doc.save(str(OUT))
print(f"Word 파일 생성 완료: {OUT}")
